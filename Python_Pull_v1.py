from docx import Document
from shotgun_api3 import Shotgun

# Setting up API for Shotgun
SERVER_PATH = ""
SCRIPT_NAME = ''
SCRIPT_KEY = ''
sg = Shotgun(SERVER_PATH, SCRIPT_NAME, SCRIPT_KEY)

# Tools used to Query the project name to their ID
proj = sg.find_one("Project", [["name", "is", "Walt Denny"]])
print proj
query = sg.schema_read()['Asset'].keys()
print query
print sg.schema_read()['Asset'].get('sg_rpm_number')

# Function to format the text that gets read by Note_Thread_Read
def formatted(x):
    x.pop("type", None)
    x.pop("id", None)
    Name = x.get('created_by', {}).get('name')
    Name2 = x.get('user', {}).get('name')
    global thread
    if 'content' in x:
        if Name == None:
            thread += "\n"
            thread += str(Name2) + ":"
            thread += "\n"
        else:
            thread += "\n"
            thread += str(Name) + ":"
            thread += "\n"
    else:
        pass
    x.pop("created_at", None)
    x.pop("user", None)
    return x

# Functions used to write out the documents to a directory in word
def writeDoc(docTitle, feed, fileName):
    # Setting up word doc for file writing
    document = Document()
    document.add_heading(fileName, 0)

    # Writing it out
    document.add_paragraph(unicode(feed, 'utf-8'))
    try :
        document.save('I:\dev_JC\_Python\Data_Pull\WD\\' + fileName + '.docx')
    except IOError:
        document.save('I:\dev_JC\_Python\Data_Pull\WD\\' + "Bad_Name" + '.docx')

# Pulls a shotgun Thread based on the id
formatting ={}
filters = [ ['id', 'is', 123]]  ##This is the controller, place whatever project ID you need docs for
project = sg.find_one('Project', filters, fields=['id'])
filters = [ ['project', 'is', {'type':'Project', 'id':project['id']}], ]
#result = sg.find('Asset', filters, fields=['sg_rpm_number', 'notes', 'id', 'sg_rpm_number'])
result = sg.find('Asset', filters, fields=['sg_rpm_number', 'cached_display_name','notes', 'id', 'sg_rpm_number'])
print result


for i in result:
    thread = ""
    dict = i['notes']
    excludeRPM = []
    excludeID = []

    # Filter out array elements that are not threads in list/bad job numbers
    if not dict:
        continue
    if i['sg_rpm_number'] in excludeRPM:
        continue

    # Pulling the thread from the dictionary
    dictStr = (dict[0]['name'])
    dictStr = dictStr.splitlines()
    try:
        thread += str(dictStr[0])
    except IndexError:
        thread += 'null'
        continue

    # Setting the fileName of the document
    try:
        fileName = 'Walt Denny - ' + i['sg_rpm_number'] + ' ' + i['cached_display_name']
    except TypeError:
        filename = 'Walt Denny - ' + i['cached_display _name']
    print fileName

    # Get's the id(TITLE) and formats it.Then appends the thread to it
    currentID = str(dict[0]['id'])
    if currentID in excludeID:
        continue
    list3 = sg.note_thread_read(int(currentID), formatting)
    for key in list3:
        thread += "\n"
        note = formatted(key)
        if 'content' in note:
            thread += str(note['content'])
        else:
            pass
    thread += '\n**END OF THREAD**'

    #Calls the Word function for file generation
    writeDoc(dictStr, thread, fileName)
